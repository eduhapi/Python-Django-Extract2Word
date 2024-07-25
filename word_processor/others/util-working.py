import openpyxl
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def extract_tables_from_excel(excel_file_path):
    wb = openpyxl.load_workbook(excel_file_path)
    extracted_tables = []
    for sheet in wb.sheetnames:
        ws = wb[sheet]
        table_data = []
        for row in ws.iter_rows(values_only=True):
            table_data.append(row)
        extracted_tables.append((sheet, table_data))
    return extracted_tables

def create_word_document(tables, output_path):
    doc = Document()
    
    # Group officers by status
    grouped_tables = {'Active': [], 'Suspended': [], 'Retired': []}
    for sheet_name, table_data in tables:
        for row in table_data[1:]:  # Skip header row
            id_, username, pno, status = row[:4]  # Assuming columns are in order
            if status.lower() in ['active', 'suspended', 'retired']:
                grouped_tables[status.capitalize()].append((id_, username, pno))
            else:
                print(f"Warning: Unexpected status '{status}' encountered.")
    
    # Create tables for each status
    for status, officer_info_list in grouped_tables.items():
        if officer_info_list:
            doc.add_heading(status, level=1)
            table = doc.add_table(rows=len(officer_info_list) + 1, cols=4)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'ID'
            hdr_cells[1].text = 'Username'
            hdr_cells[2].text = 'P/No'
            hdr_cells[3].text = 'Status'
            for i, officer_info in enumerate(officer_info_list, start=1):
                row = table.rows[i].cells
                row[0].text = str(officer_info[0])
                row[1].text = str(officer_info[1])
                row[2].text = str(officer_info[2])
                row[3].text = status

    doc.save(output_path)